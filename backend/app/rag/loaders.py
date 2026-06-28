from __future__ import annotations

from collections.abc import Iterable
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO


@dataclass(frozen=True, slots=True)
class LoadedDocument:
    text: str
    metadata: dict[str, str]


def load_policy_document(source: str | Path | bytes | BinaryIO, *, filename: str | None = None) -> LoadedDocument:
    source_name = filename or _source_name(source)
    suffix = Path(source_name).suffix.lower()
    if suffix == ".docx":
        return load_docx(source, filename=source_name)
    if suffix in {".md", ".markdown", ".txt"}:
        return load_text(source, filename=source_name)
    raise ValueError("Định dạng tài liệu chính sách chưa được hỗ trợ. Hãy dùng .docx, .md hoặc .txt.")


def load_text(source: str | Path | bytes | BinaryIO, *, filename: str | None = None) -> LoadedDocument:
    source_name = filename or _source_name(source)
    if isinstance(source, bytes):
        text = source.decode("utf-8")
    elif hasattr(source, "read"):
        data = source.read()
        text = data.decode("utf-8") if isinstance(data, bytes) else str(data)
    else:
        text = Path(source).read_text(encoding="utf-8")
    return LoadedDocument(text=text, metadata=_metadata(source_name))


def load_docx(source: str | Path | bytes | BinaryIO, *, filename: str | None = None) -> LoadedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ImportError("Cần cài python-docx để đọc file .docx") from exc

    source_name = filename or _source_name(source)
    docx_source = BytesIO(source) if isinstance(source, bytes) else source
    docx = DocxDocument(docx_source)
    blocks: list[str] = []
    for kind, item in _iter_block_items(docx):
        block = _paragraph_to_markdown(item) if kind == "paragraph" else _table_to_markdown(item)
        if block:
            blocks.append(block)
    return LoadedDocument(text="\n\n".join(blocks), metadata=_metadata(source_name))


def _paragraph_to_markdown(paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""

    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if "heading" in style_name or "title" in style_name:
        level = "1"
        for char in style_name:
            if char.isdigit():
                level = char
                break
        return f"{'#' * max(1, min(int(level), 6))} {text}"

    numbering = getattr(paragraph._p.pPr, "numPr", None) if paragraph._p.pPr is not None else None
    if numbering is not None and not text.startswith(("-", "*", "+")):
        return f"- {text}"
    return text


def _table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([" ".join(cell.text.split()) for cell in row.cells])
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    body = normalized_rows[1:]

    def render(row: list[str]) -> str:
        escaped = [cell.replace("|", "\\|") for cell in row]
        return "| " + " | ".join(escaped) + " |"

    lines = [render(header), render(["---"] * width)]
    lines.extend(render(row) for row in body)
    return "\n".join(lines)


def _iter_block_items(document) -> Iterable[tuple[str, object]]:
    from docx.document import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if not isinstance(document, DocxDocument):
        raise TypeError("Đầu vào phải là python-docx Document")

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, document)


def _metadata(source_name: str) -> dict[str, str]:
    return {
        "document_name": Path(source_name).stem,
        "source_path": source_name,
        "language": "vi",
    }


def _source_name(source: str | Path | bytes | BinaryIO) -> str:
    if isinstance(source, bytes):
        return "uploaded-policy"
    if hasattr(source, "name"):
        return str(source.name)
    return str(source)
