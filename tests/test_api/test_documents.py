import pytest
from docx import Document

from src.services.documents import reset_document_store


@pytest.fixture(autouse=True)
def clean_document_store():
    reset_document_store()
    yield
    reset_document_store()


async def _token(client, email: str, password: str) -> str:
    response = await client.post("/api/v1/auth/login", json={"email": email, "password": password})
    assert response.status_code == 200
    return response.json()["access_token"]


@pytest.mark.asyncio
async def test_admin_can_upload_and_list_document(client):
    token = await _token(client, "admin@example.com", "admin123")

    response = await client.post(
        "/api/v1/documents",
        headers={"Authorization": f"Bearer {token}"},
        json={
            "title": "Chinh sach nghi phep",
            "content": "Nhan vien co 12 ngay nghi phep nam va can bao truoc 3 ngay.",
            "visibility_roles": ["employee", "hr_admin"],
            "department_ids": [],
        },
    )

    assert response.status_code == 200
    data = response.json()
    assert data["document"]["status"] == "indexed"
    assert data["indexed_chunk_count"] >= 1
    assert "embedding" not in data["document"]

    list_response = await client.get("/api/v1/documents", headers={"Authorization": f"Bearer {token}"})
    assert list_response.status_code == 200
    assert list_response.json()["documents"][0]["title"] == "Chinh sach nghi phep"


@pytest.mark.asyncio
async def test_employee_cannot_upload_documents(client):
    token = await _token(client, "employee@example.com", "employee123")

    response = await client.post(
        "/api/v1/documents",
        headers={"Authorization": f"Bearer {token}"},
        json={
            "title": "Tai lieu noi bo",
            "content": "Noi dung",
            "visibility_roles": ["employee"],
            "department_ids": [],
        },
    )

    assert response.status_code == 403


@pytest.mark.asyncio
async def test_admin_can_delete_indexed_document_and_chunks(client):
    token = await _token(client, "admin@example.com", "admin123")
    upload = await client.post(
        "/api/v1/documents",
        headers={"Authorization": f"Bearer {token}"},
        json={
            "title": "Chinh sach nghi phep",
            "content": "Nhan vien co 12 ngay nghi phep nam va can bao truoc 3 ngay.",
            "visibility_roles": ["employee", "hr_admin"],
            "department_ids": [],
        },
    )
    assert upload.status_code == 200
    document_id = upload.json()["document"]["id"]

    delete_response = await client.delete(
        f"/api/v1/documents/{document_id}",
        headers={"Authorization": f"Bearer {token}"},
    )

    assert delete_response.status_code == 200
    assert delete_response.json()["document"]["id"] == document_id
    list_response = await client.get("/api/v1/documents", headers={"Authorization": f"Bearer {token}"})
    assert list_response.json()["documents"] == []

    employee_token = await _token(client, "employee@example.com", "employee123")
    chat_response = await client.post(
        "/api/v1/chat",
        headers={"Authorization": f"Bearer {employee_token}"},
        json={"message": "Nhan vien co bao nhieu ngay nghi phep nam?", "session_id": None},
    )
    assert chat_response.status_code == 200
    assert chat_response.json()["citations"] == []


@pytest.mark.asyncio
async def test_employee_cannot_delete_documents(client):
    admin_token = await _token(client, "admin@example.com", "admin123")
    upload = await client.post(
        "/api/v1/documents",
        headers={"Authorization": f"Bearer {admin_token}"},
        json={
            "title": "Tai lieu noi bo",
            "content": "Noi dung",
            "visibility_roles": ["employee"],
            "department_ids": [],
        },
    )
    document_id = upload.json()["document"]["id"]
    employee_token = await _token(client, "employee@example.com", "employee123")

    response = await client.delete(
        f"/api/v1/documents/{document_id}",
        headers={"Authorization": f"Bearer {employee_token}"},
    )

    assert response.status_code == 403


@pytest.mark.asyncio
async def test_documents_endpoints_appear_in_openapi(client):
    response = await client.get("/openapi.json")

    assert response.status_code == 200
    paths = response.json()["paths"]
    assert "/api/v1/documents" in paths
    assert "post" in paths["/api/v1/documents"]
    assert "get" in paths["/api/v1/documents"]
    assert "delete" in paths["/api/v1/documents/{document_id}"]
    assert "/api/v1/documents/upload" in paths


@pytest.mark.asyncio
async def test_admin_can_upload_docx_document(client, tmp_path):
    token = await _token(client, "admin@example.com", "admin123")
    docx_path = tmp_path / "so-tay.docx"
    document = Document()
    document.add_heading("CHUONG I - Chinh sach nhan su", level=1)
    document.add_heading("1.1 Nghi phep nam", level=2)
    document.add_paragraph("Nhan vien co 12 ngay nghi phep nam va can bao truoc 3 ngay.")
    document.save(docx_path)

    with docx_path.open("rb") as file:
        response = await client.post(
            "/api/v1/documents/upload",
            headers={"Authorization": f"Bearer {token}"},
            data={"title": "So tay nhan vien", "visibility_roles": ["employee", "hr_admin"]},
            files={"file": ("so-tay.docx", file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

    assert response.status_code == 200
    data = response.json()
    assert data["document"]["title"] == "So tay nhan vien"
    assert data["indexed_chunk_count"] >= 1
