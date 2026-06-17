from io import BytesIO

import pytest

from src.rag.loaders import load_policy_document
from src.services.demo_users import DEMO_USERS
from src.services.documents import create_document_from_upload, reset_document_store
from src.services.retrieval import search_policy_chunks


@pytest.fixture(autouse=True)
def clean_document_store():
    reset_document_store()
    yield
    reset_document_store()


def test_docx_loader_extracts_headings_paragraphs_and_tables():
    content = _docx_bytes()

    loaded = load_policy_document(content, filename="so-tay.docx")

    assert "# CHUONG I - Chinh sach nhan su" in loaded.text
    assert "Nhan vien co 12 ngay nghi phep nam." in loaded.text
    assert "| Che do | So ngay |" in loaded.text
    assert loaded.metadata["document_name"] == "so-tay"


def test_uploaded_docx_can_seed_retrieval_knowledge_base():
    create_document_from_upload(
        filename="so-tay.docx",
        content=_docx_bytes(),
        title="So tay nhan vien",
        visibility_roles=["employee", "hr_admin"],
    )

    citations = search_policy_chunks("Nhan vien co bao nhieu ngay nghi phep nam?", DEMO_USERS["employee@example.com"])

    assert citations
    assert citations[0].document_title == "So tay nhan vien"
    assert "12 ngay nghi phep" in citations[0].excerpt


def _docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("CHUONG I - Chinh sach nhan su", level=1)
    document.add_heading("1.1 Nghi phep nam", level=2)
    document.add_paragraph("Nhan vien co 12 ngay nghi phep nam.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Che do"
    table.cell(0, 1).text = "So ngay"
    table.cell(1, 0).text = "Nghi phep nam"
    table.cell(1, 1).text = "12"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
